from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader

from config import get_settings
from models.schemas import IngestedDocument


class UnsupportedFileError(ValueError):
    pass


class ExtractionError(RuntimeError):
    pass


class DocumentIngestionService:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}
    SUPPORTED_CONTENT_TYPES = {
        ".pdf": {"application/pdf"},
        ".docx": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
        },
        ".md": {"text/markdown", "text/plain", "application/octet-stream"},
        ".markdown": {"text/markdown", "text/plain", "application/octet-stream"},
        ".txt": {"text/plain", "application/octet-stream"},
    }

    def __init__(self) -> None:
        self.settings = get_settings()

    async def ingest(self, upload: UploadFile) -> IngestedDocument:
        if not upload.filename:
            raise UnsupportedFileError("Uploaded file must include a filename")

        extension = Path(upload.filename).suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileError(f"Unsupported file type: {extension or 'unknown'}")

        if upload.content_type and upload.content_type not in self.SUPPORTED_CONTENT_TYPES[extension]:
            raise UnsupportedFileError(f"Unsupported content type '{upload.content_type}' for {extension}")

        payload = await upload.read()
        if not payload:
            raise ExtractionError("Uploaded file is empty")
        if len(payload) > self.settings.max_upload_size_bytes:
            raise ExtractionError("Uploaded file exceeds maximum allowed size")

        text = self._extract_text(extension, payload)
        normalized = self._normalize_text(text)
        if not normalized:
            raise ExtractionError("No extractable text found in document")

        checksum = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
        return IngestedDocument(
            document_id=uuid4().hex,
            filename=upload.filename or "uploaded-document",
            file_type=extension.lstrip("."),
            checksum=checksum,
            text=normalized,
        )

    def _extract_text(self, extension: str, payload: bytes) -> str:
        if extension in {".txt", ".md", ".markdown"}:
            return payload.decode("utf-8", errors="ignore")
        if extension == ".pdf":
            return self._extract_pdf_text(payload)
        if extension == ".docx":
            return self._extract_docx_text(payload)
        raise UnsupportedFileError(f"Unsupported file type: {extension}")

    def _extract_pdf_text(self, payload: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(payload))
            texts = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(texts)
        except Exception as exc:
            raise ExtractionError("Failed to extract text from PDF") from exc

    def _extract_docx_text(self, payload: bytes) -> str:
        try:
            doc = DocxDocument(BytesIO(payload))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as exc:
            raise ExtractionError("Failed to extract text from DOCX") from exc

    def _normalize_text(self, text: str) -> str:
        lines = []
        previous_blank = False
        for raw_line in text.splitlines():
            line = " ".join(raw_line.strip().split())
            if not line:
                if not previous_blank:
                    lines.append("")
                previous_blank = True
                continue
            previous_blank = False
            lines.append(line)
        return "\n".join(lines).strip()

    def chunk_text(self, text: str) -> list[str]:
        chunk_size = self.settings.raw_chunk_size
        overlap = self.settings.raw_chunk_overlap
        if len(text) <= chunk_size:
            return [text]

        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end == len(text):
                break
            start = max(end - overlap, start + 1)
        return chunks
